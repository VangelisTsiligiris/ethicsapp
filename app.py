"""
FinTech AI Ethics & Governance Toolkit
A comprehensive tool for identifying ethical and responsible use of AI in financial technology.
Reflects latest regulations (EU AI Act, NIST AI RMF, UK FCA, Singapore MAS FEAT) and includes
risk identification, governance frameworks, and assessment tools.

Author: AI Ethics Education Initiative
Version: 1.0.0
Last Updated: November 2025
"""

import streamlit as st
import pandas as pd
import json
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Page configuration
st.set_page_config(
    page_title="FinTech AI Ethics & Governance Toolkit",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:wght@300;400;600;700&family=Fira+Code&display=swap');
    
    .main-header {
        font-family: 'Source Sans Pro', sans-serif;
        font-size: 2.5rem;
        font-weight: 700;
        color: #1a365d;
        margin-bottom: 0.5rem;
        border-bottom: 3px solid #2b6cb0;
        padding-bottom: 0.5rem;
    }
    
    .sub-header {
        font-family: 'Source Sans Pro', sans-serif;
        font-size: 1.5rem;
        font-weight: 600;
        color: #2d3748;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
    }
    
    .info-card {
        background: linear-gradient(135deg, #f7fafc 0%, #edf2f7 100%);
        border-left: 4px solid #3182ce;
        padding: 1.5rem;
        border-radius: 0 8px 8px 0;
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    .warning-card {
        background: linear-gradient(135deg, #fffaf0 0%, #feebc8 100%);
        border-left: 4px solid #dd6b20;
        padding: 1.5rem;
        border-radius: 0 8px 8px 0;
        margin: 1rem 0;
    }
    
    .success-card {
        background: linear-gradient(135deg, #f0fff4 0%, #c6f6d5 100%);
        border-left: 4px solid #38a169;
        padding: 1.5rem;
        border-radius: 0 8px 8px 0;
        margin: 1rem 0;
    }
    
    .risk-high {
        background: #fed7d7;
        color: #c53030;
        padding: 0.25rem 0.75rem;
        border-radius: 9999px;
        font-weight: 600;
        font-size: 0.875rem;
    }
    
    .risk-medium {
        background: #feebc8;
        color: #c05621;
        padding: 0.25rem 0.75rem;
        border-radius: 9999px;
        font-weight: 600;
        font-size: 0.875rem;
    }
    
    .risk-low {
        background: #c6f6d5;
        color: #276749;
        padding: 0.25rem 0.75rem;
        border-radius: 9999px;
        font-weight: 600;
        font-size: 0.875rem;
    }
    
    .metric-container {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.07);
        text-align: center;
        border: 1px solid #e2e8f0;
    }
    
    .regulation-badge {
        display: inline-block;
        padding: 0.35rem 0.75rem;
        border-radius: 6px;
        font-size: 0.8rem;
        font-weight: 600;
        margin: 0.25rem;
    }
    
    .eu-badge { background: #3182ce; color: white; }
    .us-badge { background: #805ad5; color: white; }
    .uk-badge { background: #d53f8c; color: white; }
    .sg-badge { background: #38a169; color: white; }
    
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: #f7fafc;
        border-radius: 8px 8px 0 0;
        padding: 10px 20px;
        font-weight: 600;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: #3182ce;
        color: white;
    }
    
    .framework-card {
        background: white;
        border: 1px solid #e2e8f0;
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
        transition: all 0.3s ease;
    }
    
    .framework-card:hover {
        box-shadow: 0 8px 25px rgba(0,0,0,0.1);
        transform: translateY(-2px);
    }
    
    .checklist-item {
        padding: 0.75rem;
        border-bottom: 1px solid #e2e8f0;
        display: flex;
        align-items: center;
    }
    
    .footer {
        margin-top: 3rem;
        padding: 2rem;
        background: #1a365d;
        color: white;
        border-radius: 12px;
        text-align: center;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'risk_assessment' not in st.session_state:
    st.session_state.risk_assessment = {}
if 'governance_plan' not in st.session_state:
    st.session_state.governance_plan = {}
if 'completed_assessments' not in st.session_state:
    st.session_state.completed_assessments = []

# Sidebar Navigation
st.sidebar.image("https://img.icons8.com/fluency/96/artificial-intelligence.png", width=80)
st.sidebar.markdown("## 🏛️ Navigation")

pages = {
    "🏠 Home & Overview": "home",
    "📜 Regulatory Framework": "regulations",
    "⚠️ Risk Identification Tool": "risk_tool",
    "🔧 Governance Framework Builder": "governance",
    "✅ Ethical Assessment Checklist": "assessment",
    "📊 Case Studies & Scenarios": "cases",
    "📚 Resources & Documentation": "resources",
    "💾 Export & Reports": "export"
}

selected_page = st.sidebar.radio("Select a section:", list(pages.keys()), label_visibility="collapsed")
current_page = pages[selected_page]

# ============================================
# HOME PAGE
# ============================================
if current_page == "home":
    st.markdown('<h1 class="main-header">⚖️ FinTech AI Ethics & Governance Toolkit</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-card">
        <h3>🎯 Purpose</h3>
        <p>This comprehensive toolkit helps finance professionals and students identify, assess, and govern 
        the ethical and responsible use of AI and technology in financial services. Built on the latest 
        global regulations and best practices, this tool provides practical frameworks for real-world application.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Key Statistics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        <div class="metric-container">
            <h2 style="color: #3182ce; margin: 0;">4+</h2>
            <p style="color: #718096; margin: 0;">Major Regulatory Frameworks</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="metric-container">
            <h2 style="color: #38a169; margin: 0;">50+</h2>
            <p style="color: #718096; margin: 0;">Risk Categories Covered</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="metric-container">
            <h2 style="color: #805ad5; margin: 0;">100+</h2>
            <p style="color: #718096; margin: 0;">Assessment Questions</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown("""
        <div class="metric-container">
            <h2 style="color: #dd6b20; margin: 0;">2025</h2>
            <p style="color: #718096; margin: 0;">Regulations Updated</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Core Principles Overview
    st.markdown('<h2 class="sub-header">🌟 Core Ethical Principles in FinTech AI</h2>', unsafe_allow_html=True)
    
    principles_col1, principles_col2 = st.columns(2)
    
    with principles_col1:
        st.markdown("""
        #### 🎯 Fairness
        AI systems must not discriminate against individuals or groups based on protected characteristics. 
        This includes ensuring equitable outcomes in credit scoring, lending decisions, and financial product offerings.
        
        **Key Considerations:**
        - Demographic parity in outcomes
        - Equal opportunity across groups
        - Bias detection and mitigation
        - Regular fairness audits
        
        ---
        
        #### 🔍 Transparency
        Financial institutions must be able to explain how AI systems make decisions, especially for 
        high-stakes outcomes affecting consumers.
        
        **Key Considerations:**
        - Explainable AI (XAI) implementation
        - Clear adverse action notices
        - Model documentation standards
        - Consumer-friendly explanations
        """)
    
    with principles_col2:
        st.markdown("""
        #### 👤 Accountability
        Clear lines of responsibility must exist for AI system outcomes. Organizations need defined 
        roles and escalation procedures for AI-related decisions.
        
        **Key Considerations:**
        - Senior management oversight
        - Audit trails and logging
        - Incident response procedures
        - Third-party accountability
        
        ---
        
        #### 🔒 Privacy & Security
        AI systems must protect personal data and maintain robust security measures throughout 
        the data lifecycle.
        
        **Key Considerations:**
        - Data minimization principles
        - Consent management
        - Secure data handling
        - Privacy-preserving techniques
        """)
    
    st.markdown("---")
    
    # Quick Start Guide
    st.markdown('<h2 class="sub-header">🚀 Quick Start Guide</h2>', unsafe_allow_html=True)
    
    st.markdown("""
    | Step | Action | Description |
    |------|--------|-------------|
    | 1️⃣ | **Review Regulations** | Familiarize yourself with applicable regulatory frameworks for your jurisdiction |
    | 2️⃣ | **Identify Risks** | Use the Risk Identification Tool to catalog potential AI ethics risks |
    | 3️⃣ | **Build Governance** | Create a tailored governance framework using our builder |
    | 4️⃣ | **Complete Assessment** | Run through the comprehensive ethical assessment checklist |
    | 5️⃣ | **Study Cases** | Learn from real-world scenarios and case studies |
    | 6️⃣ | **Export Report** | Generate documentation for compliance and audit purposes |
    """)

# ============================================
# REGULATORY FRAMEWORK PAGE
# ============================================
elif current_page == "regulations":
    st.markdown('<h1 class="main-header">📜 Global Regulatory Framework</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-card">
        <p>This section provides a comprehensive overview of the major AI regulations affecting 
        financial technology globally. Understanding these frameworks is essential for ensuring 
        compliance and ethical AI deployment.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Regulation Tabs
    reg_tabs = st.tabs(["🇪🇺 EU AI Act", "🇺🇸 US Frameworks", "🇬🇧 UK FCA", "🇸🇬 Singapore MAS", "📋 Comparison Matrix"])
    
    # EU AI Act Tab
    with reg_tabs[0]:
        st.markdown('<h2 class="sub-header">European Union AI Act</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <span class="regulation-badge eu-badge">Effective: August 2024 - Full Compliance: August 2026</span>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        The EU AI Act is the world's first comprehensive legal framework on artificial intelligence. 
        It establishes a risk-based approach to AI governance with significant implications for financial services.
        """)
        
        # Timeline
        st.markdown("#### 📅 Implementation Timeline")
        timeline_data = {
            "Date": ["February 2, 2025", "August 2, 2025", "August 2, 2026", "August 2, 2027"],
            "Milestone": [
                "Prohibited AI practices & AI literacy requirements take effect",
                "GPAI model rules and governance obligations apply",
                "Full application of high-risk AI system requirements",
                "Extended transition for high-risk AI in regulated products"
            ],
            "Impact": [
                "Ban on social scoring, emotion recognition in workplace, biometric categorization",
                "Transparency and copyright rules for general-purpose AI models",
                "Full compliance required for credit scoring, risk assessment AI",
                "Legacy systems in financial products must comply"
            ]
        }
        st.dataframe(pd.DataFrame(timeline_data), use_container_width=True, hide_index=True)
        
        # Risk Categories
        st.markdown("#### 🎚️ Risk Classification for FinTech")
        
        risk_col1, risk_col2 = st.columns(2)
        
        with risk_col1:
            st.markdown("""
            **🔴 High-Risk AI Systems in Finance:**
            - Credit scoring and creditworthiness evaluation
            - Life and health insurance risk assessment and pricing
            - Fraud detection systems (under certain conditions)
            - Recruitment and HR decision-making systems
            - Biometric identification systems
            
            **Requirements for High-Risk Systems:**
            - Risk management system throughout lifecycle
            - Data governance and quality requirements
            - Technical documentation
            - Record-keeping and logging
            - Transparency and user information
            - Human oversight mechanisms
            - Accuracy, robustness, and cybersecurity
            """)
        
        with risk_col2:
            st.markdown("""
            **🟡 Limited Risk AI Systems:**
            - Chatbots and conversational AI
            - Emotion recognition systems (non-prohibited uses)
            - Biometric categorization (non-prohibited uses)
            
            **Requirements:**
            - Transparency obligations (disclosure of AI interaction)
            
            ---
            
            **🟢 Minimal Risk AI Systems:**
            - AI-enabled spam filters
            - AI in video games
            - Inventory management systems
            
            **Requirements:**
            - Voluntary codes of conduct encouraged
            """)
        
        # Penalties
        st.markdown("#### ⚠️ Penalties for Non-Compliance")
        penalty_data = {
            "Violation Type": [
                "Prohibited AI practices",
                "Non-compliance with high-risk requirements",
                "Incorrect information to authorities"
            ],
            "Maximum Fine": [
                "€35 million or 7% global turnover",
                "€15 million or 3% global turnover",
                "€7.5 million or 1% global turnover"
            ]
        }
        st.dataframe(pd.DataFrame(penalty_data), use_container_width=True, hide_index=True)
    
    # US Frameworks Tab
    with reg_tabs[1]:
        st.markdown('<h2 class="sub-header">United States AI Regulatory Framework</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <span class="regulation-badge us-badge">Voluntary Framework with Sector-Specific Enforcement</span>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        The US takes a sector-specific approach to AI regulation, with multiple agencies having jurisdiction 
        over AI in financial services. Key frameworks include NIST AI RMF and agency-specific guidance.
        """)
        
        # NIST AI RMF
        st.markdown("#### 📊 NIST AI Risk Management Framework (AI RMF 1.0/2.0)")
        
        st.markdown("""
        Released January 2023, updated February 2024, with a Generative AI Profile added July 2024.
        The framework provides voluntary guidance organized around four core functions:
        """)
        
        nist_col1, nist_col2 = st.columns(2)
        
        with nist_col1:
            st.markdown("""
            **🏛️ GOVERN**
            - Establish AI governance structures
            - Define roles and responsibilities
            - Create policies and procedures
            - Foster organizational culture for responsible AI
            
            **🗺️ MAP**
            - Context and use case definition
            - Stakeholder identification
            - Risk framing and scoping
            - Impact assessment
            """)
        
        with nist_col2:
            st.markdown("""
            **📏 MEASURE**
            - Metrics and evaluation criteria
            - Testing and validation
            - Bias and fairness assessment
            - Performance monitoring
            
            **🔧 MANAGE**
            - Risk treatment and mitigation
            - Continuous monitoring
            - Incident response
            - Documentation and reporting
            """)
        
        # CFPB and Other Agencies
        st.markdown("#### 🏦 Agency-Specific Requirements")
        
        agency_data = {
            "Agency": ["CFPB", "SEC", "OCC/Federal Reserve", "FTC"],
            "Focus Area": [
                "Consumer financial protection, fair lending",
                "Investment advice, market manipulation",
                "Bank safety and soundness",
                "Consumer protection, deceptive practices"
            ],
            "Key AI Requirements": [
                "ECOA compliance, adverse action notices, no algorithmic discrimination",
                "Fiduciary duty, suitability requirements, disclosure",
                "Model risk management (SR 11-7), third-party risk",
                "Unfair/deceptive practices, algorithmic transparency"
            ],
            "Enforcement": [
                "Civil penalties, restitution, injunctive relief",
                "Civil/criminal penalties, registration revocation",
                "Cease and desist, civil money penalties",
                "Injunctions, civil penalties up to $50K/violation"
            ]
        }
        st.dataframe(pd.DataFrame(agency_data), use_container_width=True, hide_index=True)
        
        # Key CFPB Guidance
        with st.expander("📋 CFPB Key Guidance on AI"):
            st.markdown("""
            **Adverse Action Notices (2022 Circular):**
            - Creditors using complex algorithms must still provide specific reasons for adverse actions
            - "Black box" models don't exempt institutions from ECOA requirements
            - Must provide principal reasons for credit denials, even when using AI
            
            **Chatbot Guidance (2023):**
            - Chatbots cannot replace required disclosures
            - Must recognize when consumers invoke statutory rights
            - Inaccurate information may constitute UDAAP violation
            
            **Alternative Data (2024):**
            - Alternative data must meet same fair lending standards
            - Disparate impact liability applies to AI decisions
            - Regular bias testing required
            """)
    
    # UK FCA Tab
    with reg_tabs[2]:
        st.markdown('<h2 class="sub-header">UK Financial Conduct Authority (FCA)</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <span class="regulation-badge uk-badge">Principles-Based, Outcomes-Focused Regulation</span>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        The UK adopts a principles-based approach to AI regulation, relying on existing frameworks 
        while the FCA develops sector-specific guidance. The FCA confirmed in September 2025 that 
        it does not plan to introduce extra regulations for AI, instead relying on existing frameworks.
        """)
        
        # Five Principles
        st.markdown("#### 🎯 UK Government's Five AI Principles")
        
        principles_data = {
            "Principle": [
                "Safety, Security & Robustness",
                "Transparency & Explainability",
                "Fairness",
                "Accountability & Governance",
                "Contestability & Redress"
            ],
            "FCA Application": [
                "Threshold Conditions, SMCR, operational resilience requirements",
                "Consumer Duty fair value requirements, disclosure obligations",
                "Consumer Duty, Principles for Business, fair treatment of customers",
                "Senior Managers Regime, governance arrangements, SMCR",
                "Complaints handling, vulnerable customer guidance"
            ],
            "Key Requirement": [
                "Systems must be robust and secure throughout lifecycle",
                "Ability to explain AI decisions to consumers and regulators",
                "No unfair discrimination or market outcome manipulation",
                "Clear accountability at senior management level",
                "Consumers can challenge AI-driven decisions"
            ]
        }
        st.dataframe(pd.DataFrame(principles_data), use_container_width=True, hide_index=True)
        
        # FCA Initiatives
        st.markdown("#### 🔬 FCA AI Initiatives (2024-2025)")
        
        init_col1, init_col2 = st.columns(2)
        
        with init_col1:
            st.markdown("""
            **AI Lab (Launched October 2024):**
            - **Supercharged Sandbox**: Enhanced testing with NVIDIA partnership
            - **AI Live Testing**: Real-world AI deployment testing
            - **AI Spotlight**: Showcasing innovative AI solutions
            - **AI Sprint**: Stakeholder engagement events
            - **AI Input Zone**: Feedback collection mechanism
            """)
        
        with init_col2:
            st.markdown("""
            **Key Findings from 2024 Survey:**
            - 75% of regulated firms already using AI
            - Additional 10% planning AI adoption within 3 years
            - 84% have accountable person for AI framework
            - 72% report executive leadership accountability
            
            **Top Regulatory Constraints:**
            - Data protection and privacy (23% large constraint)
            - Resilience and cybersecurity (12% large constraint)
            - Consumer Duty compliance (5% large constraint)
            """)
        
        # Consumer Duty
        with st.expander("📜 Consumer Duty & AI"):
            st.markdown("""
            The Consumer Duty (effective July 2023) has significant implications for AI in financial services:
            
            **Consumer Understanding:**
            - AI-driven communications must be clear and understandable
            - Explanations of AI decisions must be accessible to consumers
            
            **Products & Services:**
            - AI-designed products must meet genuine customer needs
            - Target market assessments must account for AI-driven personalization
            
            **Price & Value:**
            - AI-driven pricing must deliver fair value
            - Dynamic pricing algorithms under scrutiny
            
            **Consumer Support:**
            - AI chatbots must provide adequate support
            - Human escalation must be available
            - Vulnerable customer needs must be recognized
            """)
    
    # Singapore MAS Tab
    with reg_tabs[3]:
        st.markdown('<h2 class="sub-header">Singapore Monetary Authority (MAS)</h2>', unsafe_allow_html=True)
        
        st.markdown("""
        <span class="regulation-badge sg-badge">FEAT Principles & Veritas Framework</span>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        Singapore has pioneered a collaborative, principles-based approach to AI governance through 
        the FEAT principles and Veritas initiative, creating practical tools for responsible AI adoption.
        """)
        
        # FEAT Principles
        st.markdown("#### 🎯 FEAT Principles (2018)")
        
        feat_col1, feat_col2 = st.columns(2)
        
        with feat_col1:
            st.markdown("""
            **F - Fairness**
            - AI systems should not systematically disadvantage individuals or groups
            - Regular assessment for unintended bias
            - Fairness metrics aligned with business context
            - Remediation processes for identified biases
            
            **E - Ethics**
            - AI use aligned with organizational values
            - Ethical review of AI applications
            - Clear ethical guidelines for AI development
            - Stakeholder impact consideration
            """)
        
        with feat_col2:
            st.markdown("""
            **A - Accountability**
            - Clear ownership and responsibility for AI systems
            - Documented decision-making processes
            - Audit trails for AI decisions
            - Escalation procedures for AI issues
            
            **T - Transparency**
            - Explainability appropriate to stakeholder needs
            - Clear communication of AI use to customers
            - Documentation of model logic and limitations
            - Disclosure of AI involvement in decisions
            """)
        
        # Veritas Initiative
        st.markdown("#### 🔧 Veritas Initiative")
        
        st.markdown("""
        The Veritas initiative provides practical tools for implementing FEAT principles:
        """)
        
        veritas_data = {
            "Phase": ["Phase 1 (2020)", "Phase 2 (2022)", "Phase 3 (2023)"],
            "Deliverables": [
                "Fairness Assessment Methodology, initial use cases",
                "Full FEAT methodologies, Veritas Toolkit v1.0, 5 white papers",
                "Veritas Toolkit v2.0, integration guidance, case studies"
            ],
            "Key Outcomes": [
                "Framework for measuring fairness in credit scoring",
                "Ethics, Accountability, Transparency methodologies added",
                "Open-source toolkit on GitHub, FI integration pilots"
            ]
        }
        st.dataframe(pd.DataFrame(veritas_data), use_container_width=True, hide_index=True)
        
        # MindForge
        with st.expander("🧠 Project MindForge - Generative AI"):
            st.markdown("""
            MindForge extends Veritas to address Generative AI risks in financial services:
            
            **Key Focus Areas:**
            - IP and copyright risks in GenAI outputs
            - Hallucination and misinformation risks
            - Data privacy in training and inference
            - Model governance for LLMs
            - Third-party AI provider oversight
            
            **Participating Organizations:**
            - Major banks: DBS, OCBC, UOB, HSBC, Standard Chartered
            - Tech providers: Google, Microsoft
            - Consulting: EY, Accenture
            
            **Recommendations:**
            - Extend FEAT principles to GenAI context
            - Implement robust testing for hallucinations
            - Establish clear accountability for AI outputs
            - Develop GenAI-specific governance frameworks
            """)
    
    # Comparison Matrix Tab
    with reg_tabs[4]:
        st.markdown('<h2 class="sub-header">Regulatory Comparison Matrix</h2>', unsafe_allow_html=True)
        
        comparison_data = {
            "Aspect": [
                "Approach",
                "Legal Status",
                "Risk Classification",
                "Explainability Requirement",
                "Bias/Fairness Testing",
                "Human Oversight",
                "Documentation",
                "Penalties",
                "Effective Date",
                "Scope"
            ],
            "EU AI Act": [
                "Prescriptive, risk-based",
                "Binding regulation",
                "4-tier (Unacceptable, High, Limited, Minimal)",
                "Mandatory for high-risk",
                "Mandatory for high-risk",
                "Mandatory for high-risk",
                "Comprehensive technical documentation",
                "Up to €35M or 7% turnover",
                "Phased 2024-2027",
                "All AI providers/deployers in EU"
            ],
            "US (NIST/Agency)": [
                "Sector-specific, voluntary framework",
                "Voluntary + agency enforcement",
                "Context-dependent",
                "ECOA requires for credit decisions",
                "Fair lending laws apply",
                "Context-dependent",
                "Model risk management (SR 11-7)",
                "Varies by agency/statute",
                "Ongoing",
                "Sector-specific jurisdiction"
            ],
            "UK FCA": [
                "Principles-based, outcomes-focused",
                "Existing rules apply",
                "Case-by-case assessment",
                "Consumer Duty requires clarity",
                "Fair treatment requirements",
                "SMCR accountability",
                "Existing governance requirements",
                "Existing FCA penalty regime",
                "Consumer Duty: July 2023",
                "FCA-regulated firms"
            ],
            "Singapore MAS": [
                "Principles-based, collaborative",
                "Voluntary guidance",
                "Context-dependent FEAT assessment",
                "FEAT Transparency principle",
                "FEAT Fairness principle",
                "FEAT Accountability principle",
                "Veritas methodology documentation",
                "Existing MAS powers",
                "2018 (FEAT), ongoing Veritas",
                "Financial institutions in Singapore"
            ]
        }
        
        st.dataframe(pd.DataFrame(comparison_data), use_container_width=True, hide_index=True)
        
        # Key Takeaways
        st.markdown("""
        <div class="success-card">
            <h4>🔑 Key Takeaways for Global Compliance</h4>
            <ul>
                <li><strong>EU AI Act sets the global standard</strong> - Many organizations are aligning with EU requirements even outside Europe</li>
                <li><strong>Existing laws still apply</strong> - No jurisdiction exempts AI from existing consumer protection, fair lending, or data protection laws</li>
                <li><strong>Documentation is universal</strong> - All frameworks require some form of documentation and audit trails</li>
                <li><strong>Human oversight is essential</strong> - All frameworks emphasize human accountability for AI decisions</li>
                <li><strong>Proactive bias testing</strong> - All frameworks expect regular assessment for discrimination and unfair outcomes</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

# ============================================
# RISK IDENTIFICATION TOOL
# ============================================
elif current_page == "risk_tool":
    st.markdown('<h1 class="main-header">⚠️ AI Risk Identification Tool</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-card">
        <p>Use this interactive tool to identify and assess AI-related risks in your FinTech application. 
        The tool categorizes risks across multiple dimensions and provides mitigation recommendations 
        aligned with global regulatory requirements.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Use Case Selection
    st.markdown("### Step 1: Define Your AI Use Case")
    
    use_case_col1, use_case_col2 = st.columns(2)
    
    with use_case_col1:
        use_case_type = st.selectbox(
            "AI Application Type:",
            [
                "Credit Scoring/Underwriting",
                "Fraud Detection",
                "Customer Service Chatbot",
                "Investment Advisory (Robo-Advisor)",
                "Anti-Money Laundering (AML)",
                "Insurance Underwriting/Pricing",
                "Marketing/Personalization",
                "Trading/Algorithmic Trading",
                "KYC/Identity Verification",
                "Collections/Debt Management",
                "Other"
            ]
        )
        
        deployment_stage = st.selectbox(
            "Deployment Stage:",
            ["Planning/Design", "Development", "Testing", "Production", "Legacy System Review"]
        )
    
    with use_case_col2:
        jurisdictions = st.multiselect(
            "Operating Jurisdictions:",
            ["European Union", "United States", "United Kingdom", "Singapore", "Other APAC", "Other"]
        )
        
        customer_type = st.multiselect(
            "Customer Types:",
            ["Retail Consumers", "Small Business", "Corporate/Institutional", "High Net Worth"]
        )
    
    st.markdown("---")
    
    # Risk Assessment Questionnaire
    st.markdown("### Step 2: Risk Assessment Questionnaire")
    
    risk_categories = {
        "Fairness & Discrimination": {
            "questions": [
                ("Does the AI system make decisions that directly impact credit access or pricing?", 3),
                ("Does the system use demographic data or proxies (zip codes, names)?", 3),
                ("Is the training data representative of all customer segments?", 2),
                ("Have you conducted disparate impact testing?", 2),
                ("Can the system's decisions be overridden by human review?", 2)
            ],
            "weight": 1.0
        },
        "Transparency & Explainability": {
            "questions": [
                ("Can you explain individual decisions to affected customers?", 3),
                ("Is the model's logic documented and understandable?", 2),
                ("Do you provide adverse action notices with specific reasons?", 3),
                ("Can regulators audit the decision-making process?", 2),
                ("Is there documentation of model limitations?", 2)
            ],
            "weight": 0.9
        },
        "Data Quality & Privacy": {
            "questions": [
                ("Is personal data collected with appropriate consent?", 3),
                ("Are data retention policies in place and enforced?", 2),
                ("Is training data checked for quality and accuracy?", 2),
                ("Are data sources from third parties properly vetted?", 2),
                ("Is data anonymization/pseudonymization used where appropriate?", 2)
            ],
            "weight": 0.85
        },
        "Security & Robustness": {
            "questions": [
                ("Is the AI system tested for adversarial attacks?", 2),
                ("Are there monitoring systems for model drift?", 2),
                ("Is the system resilient to input anomalies?", 2),
                ("Are cybersecurity measures adequate for the data sensitivity?", 3),
                ("Is there a disaster recovery plan for the AI system?", 2)
            ],
            "weight": 0.8
        },
        "Accountability & Governance": {
            "questions": [
                ("Is there a designated senior manager accountable for AI?", 3),
                ("Are there clear escalation procedures for AI issues?", 2),
                ("Is there an AI ethics committee or review board?", 2),
                ("Are third-party AI providers subject to due diligence?", 2),
                ("Is there regular board/executive reporting on AI risks?", 2)
            ],
            "weight": 0.9
        },
        "Regulatory Compliance": {
            "questions": [
                ("Have you mapped AI use to applicable regulations?", 3),
                ("Are you prepared for EU AI Act high-risk classification?", 3),
                ("Is there a process for regulatory change monitoring?", 2),
                ("Have you conducted a compliance gap analysis?", 2),
                ("Is regulatory reporting capability in place?", 2)
            ],
            "weight": 1.0
        }
    }
    
    # Store responses
    responses = {}
    risk_scores = {}
    
    for category, config in risk_categories.items():
        with st.expander(f"📋 {category}", expanded=False):
            responses[category] = {}
            category_score = 0
            max_score = 0
            
            for i, (question, weight) in enumerate(config["questions"]):
                response = st.radio(
                    question,
                    ["Yes - Fully Implemented", "Partial - In Progress", "No - Not Addressed", "N/A"],
                    key=f"{category}_{i}",
                    horizontal=True
                )
                responses[category][question] = response
                
                # Calculate score
                max_score += weight
                if response == "Yes - Fully Implemented":
                    category_score += weight
                elif response == "Partial - In Progress":
                    category_score += weight * 0.5
                elif response == "N/A":
                    max_score -= weight  # Don't count N/A against score
            
            if max_score > 0:
                risk_scores[category] = {
                    "score": (category_score / max_score) * 100,
                    "weight": config["weight"]
                }
            else:
                risk_scores[category] = {"score": 100, "weight": config["weight"]}
    
    st.markdown("---")
    
    # Risk Analysis Results
    if st.button("🔍 Generate Risk Analysis", type="primary"):
        st.markdown("### Step 3: Risk Analysis Results")
        
        # Calculate overall risk
        total_weighted_score = sum(s["score"] * s["weight"] for s in risk_scores.values())
        total_weight = sum(s["weight"] for s in risk_scores.values())
        overall_score = total_weighted_score / total_weight if total_weight > 0 else 0
        
        # Determine risk level
        if overall_score >= 80:
            risk_level = "Low"
            risk_color = "#38a169"
            risk_class = "risk-low"
        elif overall_score >= 60:
            risk_level = "Medium"
            risk_color = "#dd6b20"
            risk_class = "risk-medium"
        else:
            risk_level = "High"
            risk_color = "#c53030"
            risk_class = "risk-high"
        
        # Display overall score
        score_col1, score_col2, score_col3 = st.columns([1, 2, 1])
        
        with score_col2:
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=overall_score,
                domain={'x': [0, 1], 'y': [0, 1]},
                title={'text': "Overall Risk Readiness Score"},
                gauge={
                    'axis': {'range': [0, 100]},
                    'bar': {'color': risk_color},
                    'steps': [
                        {'range': [0, 60], 'color': "#fed7d7"},
                        {'range': [60, 80], 'color': "#feebc8"},
                        {'range': [80, 100], 'color': "#c6f6d5"}
                    ],
                    'threshold': {
                        'line': {'color': "black", 'width': 4},
                        'thickness': 0.75,
                        'value': overall_score
                    }
                }
            ))
            fig.update_layout(height=300)
            st.plotly_chart(fig, use_container_width=True)
        
        st.markdown(f"""
        <div style="text-align: center; margin: 1rem 0;">
            <span class="{risk_class}" style="font-size: 1.25rem;">Risk Level: {risk_level}</span>
        </div>
        """, unsafe_allow_html=True)
        
        # Category breakdown
        st.markdown("#### Category Breakdown")
        
        category_scores = pd.DataFrame([
            {"Category": cat, "Score": data["score"], "Risk Level": 
             "Low" if data["score"] >= 80 else "Medium" if data["score"] >= 60 else "High"}
            for cat, data in risk_scores.items()
        ])
        
        fig2 = px.bar(
            category_scores,
            x="Category",
            y="Score",
            color="Risk Level",
            color_discrete_map={"Low": "#38a169", "Medium": "#dd6b20", "High": "#c53030"},
            title="Risk Readiness by Category"
        )
        fig2.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig2, use_container_width=True)
        
        # High-Risk Jurisdiction Alert
        if "European Union" in jurisdictions:
            if use_case_type in ["Credit Scoring/Underwriting", "Insurance Underwriting/Pricing"]:
                st.markdown("""
                <div class="warning-card">
                    <h4>⚠️ EU AI Act High-Risk Classification Alert</h4>
                    <p>Your use case (<strong>{}</strong>) is likely classified as <strong>HIGH-RISK</strong> 
                    under the EU AI Act. This requires:</p>
                    <ul>
                        <li>Comprehensive risk management system</li>
                        <li>Data governance and quality requirements</li>
                        <li>Technical documentation</li>
                        <li>Automatic logging and record-keeping</li>
                        <li>Transparency and user information</li>
                        <li>Human oversight mechanisms</li>
                        <li>Accuracy, robustness, and cybersecurity measures</li>
                    </ul>
                    <p><strong>Deadline:</strong> Full compliance required by August 2, 2026</p>
                </div>
                """.format(use_case_type), unsafe_allow_html=True)
        
        # Mitigation Recommendations
        st.markdown("#### 🛠️ Priority Mitigation Recommendations")
        
        low_score_categories = [cat for cat, data in risk_scores.items() if data["score"] < 70]
        
        mitigation_recommendations = {
            "Fairness & Discrimination": [
                "Implement demographic parity testing across all protected characteristics",
                "Document and justify any use of proxy variables",
                "Establish regular fairness audits (quarterly minimum)",
                "Create human override procedures for edge cases",
                "Develop bias incident response procedures"
            ],
            "Transparency & Explainability": [
                "Implement LIME or SHAP for individual decision explanations",
                "Create consumer-friendly explanation templates for adverse actions",
                "Document model logic, features, and limitations",
                "Establish regulatory inquiry response procedures",
                "Train customer service on explaining AI decisions"
            ],
            "Data Quality & Privacy": [
                "Conduct comprehensive data mapping and consent audit",
                "Implement data quality checks in AI pipeline",
                "Review and update data retention policies",
                "Conduct third-party data source due diligence",
                "Implement privacy-enhancing technologies where feasible"
            ],
            "Security & Robustness": [
                "Conduct adversarial robustness testing",
                "Implement model drift monitoring",
                "Establish input validation and anomaly detection",
                "Review cybersecurity measures for AI infrastructure",
                "Develop AI-specific incident response procedures"
            ],
            "Accountability & Governance": [
                "Designate senior AI accountability officer",
                "Establish AI ethics committee/review board",
                "Create third-party AI provider governance framework",
                "Implement regular board reporting on AI risks",
                "Develop clear escalation procedures"
            ],
            "Regulatory Compliance": [
                "Conduct comprehensive regulatory mapping exercise",
                "Assess EU AI Act classification and requirements",
                "Establish regulatory change monitoring process",
                "Conduct gap analysis against all applicable frameworks",
                "Develop regulatory reporting capabilities"
            ]
        }
        
        if low_score_categories:
            for category in low_score_categories:
                with st.expander(f"🔧 {category} - Recommendations"):
                    for i, rec in enumerate(mitigation_recommendations.get(category, []), 1):
                        st.markdown(f"{i}. {rec}")
        else:
            st.markdown("""
            <div class="success-card">
                <h4>✅ Good Risk Posture</h4>
                <p>Your assessment indicates a generally strong risk management posture. 
                Continue to monitor and maintain your controls, and conduct regular reviews 
                as regulations evolve.</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Save to session state
        st.session_state.risk_assessment = {
            "timestamp": datetime.now().isoformat(),
            "use_case": use_case_type,
            "jurisdictions": jurisdictions,
            "overall_score": overall_score,
            "risk_level": risk_level,
            "category_scores": {cat: data["score"] for cat, data in risk_scores.items()},
            "responses": responses
        }

# ============================================
# GOVERNANCE FRAMEWORK BUILDER
# ============================================
elif current_page == "governance":
    st.markdown('<h1 class="main-header">🔧 Governance Framework Builder</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-card">
        <p>Build a customized AI governance framework for your organization. This tool helps you 
        establish the policies, procedures, and controls needed for responsible AI deployment in 
        financial services.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Organization Profile
    st.markdown("### 📋 Organization Profile")
    
    org_col1, org_col2 = st.columns(2)
    
    with org_col1:
        org_size = st.selectbox(
            "Organization Size:",
            ["Startup (<50 employees)", "SME (50-250 employees)", 
             "Mid-size (250-1000 employees)", "Large Enterprise (1000+ employees)"]
        )
        
        primary_business = st.selectbox(
            "Primary Business:",
            ["Retail Banking", "Commercial Banking", "Investment Banking",
             "Asset Management", "Insurance", "Payments/FinTech", 
             "Lending Platform", "WealthTech", "RegTech", "Other"]
        )
    
    with org_col2:
        regulatory_status = st.selectbox(
            "Regulatory Status:",
            ["Fully Regulated (Bank License)", "Regulated (Other License)",
             "Registered/Authorized", "Unregulated (Partner with Licensed Entity)",
             "Seeking Authorization"]
        )
        
        ai_maturity = st.selectbox(
            "AI Maturity Level:",
            ["Exploring (No Production AI)", "Emerging (1-2 Production Systems)",
             "Established (Multiple Production Systems)", "Advanced (AI-First Organization)"]
        )
    
    st.markdown("---")
    
    # Framework Components
    st.markdown("### 🏗️ Governance Framework Components")
    
    framework_tabs = st.tabs([
        "1️⃣ Governance Structure", 
        "2️⃣ Policies & Procedures",
        "3️⃣ Risk Management",
        "4️⃣ Lifecycle Controls",
        "5️⃣ Monitoring & Reporting"
    ])
    
    governance_plan = {}
    
    # Tab 1: Governance Structure
    with framework_tabs[0]:
        st.markdown("#### Organizational Structure for AI Governance")
        
        st.markdown("""
        Define the key roles and committees responsible for AI governance:
        """)
        
        # Accountability Framework
        st.markdown("##### 👤 Accountability Framework")
        
        acc_col1, acc_col2 = st.columns(2)
        
        with acc_col1:
            ai_officer = st.text_input("Chief AI/Data Officer (or equivalent):", 
                                        placeholder="Name and title")
            ai_sponsor = st.text_input("Executive Sponsor for AI:", 
                                        placeholder="Board/C-suite member")
        
        with acc_col2:
            ai_risk_owner = st.text_input("AI Risk Owner:", 
                                           placeholder="Head of Risk/CRO")
            ai_ethics_owner = st.text_input("AI Ethics/Compliance Owner:", 
                                             placeholder="CCO/Ethics Officer")
        
        # Committees
        st.markdown("##### 🏛️ Committee Structure")
        
        has_ai_committee = st.checkbox("AI Ethics/Governance Committee established")
        if has_ai_committee:
            committee_frequency = st.selectbox(
                "Committee Meeting Frequency:",
                ["Monthly", "Bi-monthly", "Quarterly"]
            )
            committee_chair = st.text_input("Committee Chair:")
            committee_members = st.text_area("Committee Members (one per line):")
        
        has_model_committee = st.checkbox("Model Risk Committee established")
        has_data_committee = st.checkbox("Data Governance Committee established")
        
        # Lines of Defense
        st.markdown("##### 🛡️ Three Lines of Defense")
        
        lod_col1, lod_col2, lod_col3 = st.columns(3)
        
        with lod_col1:
            st.markdown("""
            **1st Line: Business Units**
            """)
            first_line_resp = st.multiselect(
                "Responsibilities:",
                [
                    "AI use case identification",
                    "Initial risk assessment",
                    "Model performance monitoring",
                    "User acceptance testing",
                    "Incident escalation"
                ],
                default=["AI use case identification", "Initial risk assessment"]
            )
        
        with lod_col2:
            st.markdown("""
            **2nd Line: Risk & Compliance**
            """)
            second_line_resp = st.multiselect(
                "Responsibilities:",
                [
                    "Independent model validation",
                    "Regulatory compliance review",
                    "Fairness/bias testing",
                    "Policy development",
                    "Training and awareness"
                ],
                default=["Independent model validation", "Regulatory compliance review"]
            )
        
        with lod_col3:
            st.markdown("""
            **3rd Line: Internal Audit**
            """)
            third_line_resp = st.multiselect(
                "Responsibilities:",
                [
                    "Governance effectiveness audit",
                    "Process compliance audit",
                    "Controls testing",
                    "Third-party audit oversight",
                    "Board reporting"
                ],
                default=["Governance effectiveness audit", "Process compliance audit"]
            )
        
        governance_plan["structure"] = {
            "ai_officer": ai_officer,
            "ai_sponsor": ai_sponsor,
            "ai_risk_owner": ai_risk_owner,
            "ai_ethics_owner": ai_ethics_owner,
            "has_ai_committee": has_ai_committee,
            "has_model_committee": has_model_committee,
            "has_data_committee": has_data_committee,
            "first_line": first_line_resp,
            "second_line": second_line_resp,
            "third_line": third_line_resp
        }
    
    # Tab 2: Policies & Procedures
    with framework_tabs[1]:
        st.markdown("#### Core AI Policies")
        
        st.markdown("""
        Select the policies your organization has or needs to develop:
        """)
        
        policy_status = {}
        
        core_policies = [
            ("AI Ethics Policy", "Defines ethical principles and values for AI use"),
            ("AI Risk Management Policy", "Framework for identifying and managing AI risks"),
            ("Model Risk Management Policy", "SR 11-7 aligned model governance"),
            ("Data Governance Policy", "Data quality, privacy, and usage standards"),
            ("Third-Party AI Policy", "Vendor due diligence and oversight"),
            ("AI Transparency Policy", "Explainability and disclosure requirements"),
            ("AI Fairness Policy", "Bias prevention and fairness testing"),
            ("AI Incident Management Policy", "Response procedures for AI failures")
        ]
        
        for policy, description in core_policies:
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"**{policy}**")
                st.caption(description)
            with col2:
                policy_status[policy] = st.selectbox(
                    f"Status",
                    ["Not Started", "In Development", "Under Review", "Approved", "N/A"],
                    key=f"policy_{policy}"
                )
        
        governance_plan["policies"] = policy_status
        
        # Procedures
        st.markdown("---")
        st.markdown("#### Key Procedures")
        
        procedure_status = {}
        
        key_procedures = [
            "AI Use Case Approval Process",
            "Model Validation Procedures",
            "Fairness Testing Procedures",
            "AI Incident Escalation Procedures",
            "AI Change Management Procedures",
            "Third-Party AI Due Diligence Procedures",
            "AI Documentation Standards",
            "Consumer Disclosure Procedures"
        ]
        
        for procedure in key_procedures:
            col1, col2 = st.columns([3, 1])
            with col1:
                st.markdown(f"📄 {procedure}")
            with col2:
                procedure_status[procedure] = st.selectbox(
                    f"Status",
                    ["Not Started", "In Development", "Under Review", "Approved", "N/A"],
                    key=f"proc_{procedure}"
                )
        
        governance_plan["procedures"] = procedure_status
    
    # Tab 3: Risk Management
    with framework_tabs[2]:
        st.markdown("#### AI Risk Management Framework")
        
        st.markdown("""
        Define your approach to AI risk identification, assessment, and management:
        """)
        
        # Risk Taxonomy
        st.markdown("##### 📊 Risk Taxonomy")
        
        risk_taxonomy = st.multiselect(
            "Risk Categories to Address:",
            [
                "Model Risk (performance, drift, failure)",
                "Fairness/Discrimination Risk",
                "Data Quality Risk",
                "Privacy/Data Protection Risk",
                "Cybersecurity Risk",
                "Operational/Reliability Risk",
                "Regulatory/Compliance Risk",
                "Reputational Risk",
                "Third-Party/Concentration Risk",
                "Strategic Risk (misalignment with objectives)"
            ],
            default=[
                "Model Risk (performance, drift, failure)",
                "Fairness/Discrimination Risk",
                "Regulatory/Compliance Risk"
            ]
        )
        
        # Risk Assessment Approach
        st.markdown("##### 🎯 Risk Assessment Methodology")
        
        risk_approach = st.radio(
            "Risk Assessment Approach:",
            [
                "Quantitative (numerical scoring, statistical analysis)",
                "Qualitative (expert judgment, categorical ratings)",
                "Hybrid (combination of quantitative and qualitative)"
            ]
        )
        
        risk_frequency = st.selectbox(
            "Risk Assessment Frequency:",
            ["Continuous/Real-time", "Monthly", "Quarterly", "Annually", "Event-driven only"]
        )
        
        # Risk Appetite
        st.markdown("##### 🎚️ Risk Appetite")
        
        st.markdown("Define your organization's AI risk appetite:")
        
        appetite_col1, appetite_col2 = st.columns(2)
        
        with appetite_col1:
            fairness_appetite = st.select_slider(
                "Fairness Risk Tolerance:",
                options=["Zero Tolerance", "Very Low", "Low", "Moderate", "High"]
            )
            model_risk_appetite = st.select_slider(
                "Model Risk Tolerance:",
                options=["Zero Tolerance", "Very Low", "Low", "Moderate", "High"]
            )
        
        with appetite_col2:
            regulatory_appetite = st.select_slider(
                "Regulatory Risk Tolerance:",
                options=["Zero Tolerance", "Very Low", "Low", "Moderate", "High"]
            )
            reputational_appetite = st.select_slider(
                "Reputational Risk Tolerance:",
                options=["Zero Tolerance", "Very Low", "Low", "Moderate", "High"]
            )
        
        governance_plan["risk_management"] = {
            "taxonomy": risk_taxonomy,
            "approach": risk_approach,
            "frequency": risk_frequency,
            "appetite": {
                "fairness": fairness_appetite,
                "model": model_risk_appetite,
                "regulatory": regulatory_appetite,
                "reputational": reputational_appetite
            }
        }
    
    # Tab 4: Lifecycle Controls
    with framework_tabs[3]:
        st.markdown("#### AI System Lifecycle Controls")
        
        st.markdown("""
        Define controls at each stage of the AI system lifecycle:
        """)
        
        lifecycle_stages = {
            "🎯 Planning & Design": [
                "Use case business justification",
                "Ethical impact assessment",
                "Regulatory classification",
                "Data requirements analysis",
                "Stakeholder identification",
                "Success criteria definition"
            ],
            "🔨 Development": [
                "Data quality validation",
                "Feature selection review",
                "Model selection justification",
                "Bias testing during training",
                "Code review requirements",
                "Version control"
            ],
            "🧪 Testing & Validation": [
                "Independent model validation",
                "Fairness/bias testing",
                "Performance testing",
                "Stress testing",
                "User acceptance testing",
                "Security testing"
            ],
            "🚀 Deployment": [
                "Deployment approval process",
                "Rollback procedures",
                "Monitoring implementation",
                "Documentation completeness check",
                "Training completion verification",
                "Regulatory notification (if required)"
            ],
            "📊 Monitoring & Maintenance": [
                "Performance monitoring",
                "Drift detection",
                "Fairness monitoring",
                "Incident logging",
                "Periodic revalidation",
                "Decommissioning criteria"
            ]
        }
        
        lifecycle_controls = {}
        
        for stage, controls in lifecycle_stages.items():
            st.markdown(f"##### {stage}")
            lifecycle_controls[stage] = st.multiselect(
                f"Select required controls:",
                controls,
                default=controls[:3],
                key=f"lifecycle_{stage}"
            )
        
        governance_plan["lifecycle_controls"] = lifecycle_controls
    
    # Tab 5: Monitoring & Reporting
    with framework_tabs[4]:
        st.markdown("#### Monitoring & Reporting Framework")
        
        # Key Metrics
        st.markdown("##### 📈 Key Performance Indicators (KPIs)")
        
        kpi_options = [
            "Model accuracy/performance metrics",
            "Fairness metrics (demographic parity, equal opportunity)",
            "Drift detection scores",
            "Incident count and severity",
            "Regulatory findings",
            "Customer complaints related to AI",
            "Adverse action appeal rates",
            "Third-party SLA compliance",
            "Model validation coverage",
            "Training completion rates"
        ]
        
        selected_kpis = st.multiselect(
            "Select KPIs to Track:",
            kpi_options,
            default=kpi_options[:5]
        )
        
        # Reporting Structure
        st.markdown("##### 📋 Reporting Structure")
        
        report_col1, report_col2 = st.columns(2)
        
        with report_col1:
            st.markdown("**Board Reporting:**")
            board_frequency = st.selectbox(
                "Frequency:",
                ["Quarterly", "Semi-annually", "Annually"],
                key="board_freq"
            )
            board_content = st.multiselect(
                "Content:",
                ["AI risk dashboard", "Regulatory compliance status", 
                 "Incident summary", "Strategic initiatives", "External audit findings"],
                default=["AI risk dashboard", "Regulatory compliance status"]
            )
        
        with report_col2:
            st.markdown("**Management Reporting:**")
            mgmt_frequency = st.selectbox(
                "Frequency:",
                ["Weekly", "Monthly", "Quarterly"],
                key="mgmt_freq"
            )
            mgmt_content = st.multiselect(
                "Content:",
                ["Operational metrics", "Issue/incident log", "Project status",
                 "Validation results", "Regulatory updates"],
                default=["Operational metrics", "Issue/incident log"]
            )
        
        # Audit Requirements
        st.markdown("##### 🔍 Audit Requirements")
        
        internal_audit = st.checkbox("Annual internal audit of AI governance", value=True)
        external_audit = st.checkbox("External/independent AI audit")
        regulatory_exam = st.checkbox("Regulatory examination preparation")
        
        governance_plan["monitoring"] = {
            "kpis": selected_kpis,
            "board_reporting": {"frequency": board_frequency, "content": board_content},
            "management_reporting": {"frequency": mgmt_frequency, "content": mgmt_content},
            "internal_audit": internal_audit,
            "external_audit": external_audit,
            "regulatory_exam": regulatory_exam
        }
    
    st.markdown("---")
    
    # Generate Framework
    if st.button("📄 Generate Governance Framework", type="primary"):
        st.session_state.governance_plan = governance_plan
        
        st.markdown("### ✅ Governance Framework Generated")
        
        st.markdown("""
        <div class="success-card">
            <h4>Framework Summary</h4>
            <p>Your customized AI governance framework has been generated. 
            You can export this framework from the Export & Reports section.</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Summary Statistics
        summary_col1, summary_col2, summary_col3 = st.columns(3)
        
        with summary_col1:
            policies_defined = len([p for p, s in governance_plan.get("policies", {}).items() 
                                   if s not in ["Not Started", "N/A"]])
            st.metric("Policies Defined", f"{policies_defined}/8")
        
        with summary_col2:
            procedures_defined = len([p for p, s in governance_plan.get("procedures", {}).items() 
                                     if s not in ["Not Started", "N/A"]])
            st.metric("Procedures Defined", f"{procedures_defined}/8")
        
        with summary_col3:
            total_controls = sum(len(controls) for controls in 
                               governance_plan.get("lifecycle_controls", {}).values())
            st.metric("Lifecycle Controls", total_controls)

# ============================================
# ETHICAL ASSESSMENT CHECKLIST
# ============================================
elif current_page == "assessment":
    st.markdown('<h1 class="main-header">✅ Ethical Assessment Checklist</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-card">
        <p>Complete this comprehensive checklist to assess the ethical readiness of your AI system. 
        This checklist is aligned with major regulatory frameworks including the EU AI Act, 
        NIST AI RMF, UK FCA guidance, and Singapore MAS FEAT principles.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # System Information
    st.markdown("### 📝 System Information")
    
    sys_col1, sys_col2 = st.columns(2)
    
    with sys_col1:
        system_name = st.text_input("AI System Name:", placeholder="e.g., Credit Decision Engine v2.0")
        system_owner = st.text_input("System Owner:", placeholder="Name and department")
    
    with sys_col2:
        assessment_date = st.date_input("Assessment Date:", value=datetime.now())
        assessor_name = st.text_input("Assessor:", placeholder="Your name")
    
    st.markdown("---")
    
    # Comprehensive Checklist
    assessment_sections = {
        "1. Fairness & Non-Discrimination": {
            "description": "Ensure the AI system does not discriminate against protected groups",
            "items": [
                ("1.1", "Protected characteristics (race, gender, age, etc.) are not used as direct inputs", "Critical"),
                ("1.2", "Proxy variables have been analyzed for correlation with protected characteristics", "Critical"),
                ("1.3", "Training data has been assessed for representation bias", "Critical"),
                ("1.4", "Disparate impact testing has been conducted", "Critical"),
                ("1.5", "Fairness metrics (demographic parity, equal opportunity) are monitored", "High"),
                ("1.6", "Adverse impact remediation procedures are documented", "High"),
                ("1.7", "Human override is available for edge cases", "High"),
                ("1.8", "Regular fairness audits are scheduled (at least annually)", "Medium")
            ]
        },
        "2. Transparency & Explainability": {
            "description": "Ensure decisions can be explained and understood",
            "items": [
                ("2.1", "Individual decisions can be explained to affected consumers", "Critical"),
                ("2.2", "Adverse action notices include specific, accurate reasons", "Critical"),
                ("2.3", "Model logic and key features are documented", "High"),
                ("2.4", "Explainability tools (LIME, SHAP) are implemented where appropriate", "High"),
                ("2.5", "Consumer-friendly explanations are available", "High"),
                ("2.6", "Regulators can audit the decision-making process", "Critical"),
                ("2.7", "Model limitations are documented and communicated", "Medium"),
                ("2.8", "Technical documentation meets EU AI Act standards (if applicable)", "High")
            ]
        },
        "3. Accountability & Governance": {
            "description": "Establish clear lines of responsibility",
            "items": [
                ("3.1", "A senior manager is designated as accountable for the AI system", "Critical"),
                ("3.2", "Roles and responsibilities are clearly defined and documented", "High"),
                ("3.3", "Escalation procedures for AI issues are established", "High"),
                ("3.4", "An AI ethics/governance committee reviews high-risk systems", "Medium"),
                ("3.5", "Third-party AI providers are subject to due diligence", "High"),
                ("3.6", "Contracts with AI vendors include appropriate liability provisions", "Medium"),
                ("3.7", "Regular board/executive reporting on AI risks is in place", "Medium"),
                ("3.8", "Audit trails capture all AI decisions and can be retrieved", "Critical")
            ]
        },
        "4. Data Quality & Privacy": {
            "description": "Ensure data is accurate, relevant, and protected",
            "items": [
                ("4.1", "Personal data is collected with appropriate legal basis/consent", "Critical"),
                ("4.2", "Data minimization principles are applied", "High"),
                ("4.3", "Training data quality has been validated", "High"),
                ("4.4", "Data sources are documented and vetted", "High"),
                ("4.5", "Data retention policies are defined and enforced", "High"),
                ("4.6", "Data subject rights (access, deletion) can be fulfilled", "Critical"),
                ("4.7", "Privacy impact assessment has been conducted", "High"),
                ("4.8", "Cross-border data transfer requirements are met", "High")
            ]
        },
        "5. Security & Robustness": {
            "description": "Ensure the AI system is secure and reliable",
            "items": [
                ("5.1", "The AI system has been tested for adversarial attacks", "High"),
                ("5.2", "Input validation and anomaly detection are implemented", "High"),
                ("5.3", "Model drift monitoring is in place", "High"),
                ("5.4", "Cybersecurity measures are appropriate for data sensitivity", "Critical"),
                ("5.5", "Disaster recovery/business continuity plans include AI systems", "High"),
                ("5.6", "Access controls limit who can modify the AI system", "High"),
                ("5.7", "Model versioning and rollback capabilities exist", "Medium"),
                ("5.8", "Stress testing under extreme conditions has been performed", "Medium")
            ]
        },
        "6. Human Oversight": {
            "description": "Maintain appropriate human control over AI decisions",
            "items": [
                ("6.1", "Human review is required for high-stakes decisions", "Critical"),
                ("6.2", "Staff can understand and challenge AI recommendations", "High"),
                ("6.3", "Override mechanisms are available and documented", "High"),
                ("6.4", "Staff receive training on AI system use and limitations", "High"),
                ("6.5", "Escalation paths for uncertain cases are defined", "Medium"),
                ("6.6", "Human reviewers have sufficient time and information", "Medium"),
                ("6.7", "Override decisions are logged and analyzed", "Medium"),
                ("6.8", "Alert thresholds trigger human review appropriately", "High")
            ]
        },
        "7. Consumer Protection": {
            "description": "Protect consumer rights and interests",
            "items": [
                ("7.1", "Consumers are informed when AI is used in decisions affecting them", "High"),
                ("7.2", "Complaint and appeal procedures are accessible", "Critical"),
                ("7.3", "Vulnerable consumers are identified and protected", "High"),
                ("7.4", "AI-driven products meet genuine customer needs", "High"),
                ("7.5", "Pricing decisions are fair and non-exploitative", "High"),
                ("7.6", "Marketing personalization respects consumer preferences", "Medium"),
                ("7.7", "Consumer support can address AI-related queries", "High"),
                ("7.8", "Redress mechanisms are available for AI errors", "Critical")
            ]
        },
        "8. Regulatory Compliance": {
            "description": "Meet all applicable regulatory requirements",
            "items": [
                ("8.1", "All applicable regulations have been identified and mapped", "Critical"),
                ("8.2", "EU AI Act classification and requirements are addressed (if applicable)", "Critical"),
                ("8.3", "Fair lending/ECOA requirements are met (if applicable)", "Critical"),
                ("8.4", "Sector-specific regulations are addressed", "High"),
                ("8.5", "Regulatory change monitoring process is in place", "High"),
                ("8.6", "Regulatory reporting capabilities are established", "High"),
                ("8.7", "Regulatory sandbox participation considered (if appropriate)", "Low"),
                ("8.8", "Legal review of AI system compliance has been conducted", "High")
            ]
        }
    }
    
    # Assessment Form
    assessment_results = {}
    section_scores = {}
    
    for section, content in assessment_sections.items():
        with st.expander(f"📋 {section}", expanded=False):
            st.markdown(f"*{content['description']}*")
            
            section_results = {}
            compliant_count = 0
            total_weighted = 0
            compliant_weighted = 0
            
            for item_id, item_text, priority in content["items"]:
                priority_weight = {"Critical": 3, "High": 2, "Medium": 1, "Low": 0.5}.get(priority, 1)
                total_weighted += priority_weight
                
                col1, col2, col3 = st.columns([0.5, 3, 1.5])
                
                with col1:
                    st.markdown(f"**{item_id}**")
                
                with col2:
                    st.markdown(item_text)
                    if priority == "Critical":
                        st.markdown("🔴 *Critical*")
                    elif priority == "High":
                        st.markdown("🟠 *High*")
                    else:
                        st.markdown("🟡 *Medium*")
                
                with col3:
                    status = st.selectbox(
                        "Status",
                        ["Not Assessed", "Compliant", "Partial", "Non-Compliant", "N/A"],
                        key=f"check_{item_id}",
                        label_visibility="collapsed"
                    )
                    section_results[item_id] = {"text": item_text, "priority": priority, "status": status}
                    
                    if status == "Compliant":
                        compliant_weighted += priority_weight
                    elif status == "Partial":
                        compliant_weighted += priority_weight * 0.5
                    elif status == "N/A":
                        total_weighted -= priority_weight
            
            assessment_results[section] = section_results
            if total_weighted > 0:
                section_scores[section] = (compliant_weighted / total_weighted) * 100
            else:
                section_scores[section] = 100
    
    st.markdown("---")
    
    # Generate Assessment Report
    if st.button("📊 Generate Assessment Report", type="primary"):
        st.markdown("### 📊 Assessment Results")
        
        # Overall Score
        overall_score = sum(section_scores.values()) / len(section_scores) if section_scores else 0
        
        if overall_score >= 80:
            overall_status = "Ready for Production"
            status_color = "#38a169"
        elif overall_score >= 60:
            overall_status = "Needs Improvement"
            status_color = "#dd6b20"
        else:
            overall_status = "Not Ready"
            status_color = "#c53030"
        
        # Summary Cards
        score_col1, score_col2, score_col3 = st.columns(3)
        
        with score_col1:
            st.markdown(f"""
            <div class="metric-container">
                <h2 style="color: {status_color}; margin: 0;">{overall_score:.1f}%</h2>
                <p style="color: #718096; margin: 0;">Overall Compliance Score</p>
            </div>
            """, unsafe_allow_html=True)
        
        with score_col2:
            st.markdown(f"""
            <div class="metric-container">
                <h2 style="color: {status_color}; margin: 0;">{overall_status}</h2>
                <p style="color: #718096; margin: 0;">Readiness Status</p>
            </div>
            """, unsafe_allow_html=True)
        
        with score_col3:
            critical_issues = sum(
                1 for section in assessment_results.values() 
                for item in section.values() 
                if item["priority"] == "Critical" and item["status"] == "Non-Compliant"
            )
            st.markdown(f"""
            <div class="metric-container">
                <h2 style="color: {'#c53030' if critical_issues > 0 else '#38a169'}; margin: 0;">{critical_issues}</h2>
                <p style="color: #718096; margin: 0;">Critical Issues</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Section Breakdown
        st.markdown("#### Section Breakdown")
        
        section_df = pd.DataFrame([
            {"Section": section.split(". ")[1], "Score": f"{score:.1f}%", 
             "Status": "✅ Pass" if score >= 80 else "⚠️ Review" if score >= 60 else "❌ Fail"}
            for section, score in section_scores.items()
        ])
        
        fig = px.bar(
            section_df,
            x="Section",
            y=[float(s.replace("%", "")) for s in section_df["Score"]],
            color=[float(s.replace("%", "")) for s in section_df["Score"]],
            color_continuous_scale=["#c53030", "#dd6b20", "#38a169"],
            title="Compliance Score by Section"
        )
        fig.update_layout(xaxis_tickangle=-45, showlegend=False)
        fig.update_coloraxes(showscale=False)
        st.plotly_chart(fig, use_container_width=True)
        
        # Non-Compliant Items
        st.markdown("#### ⚠️ Items Requiring Attention")
        
        attention_items = []
        for section, items in assessment_results.items():
            for item_id, item_data in items.items():
                if item_data["status"] in ["Non-Compliant", "Partial"]:
                    attention_items.append({
                        "ID": item_id,
                        "Section": section.split(". ")[1],
                        "Item": item_data["text"][:50] + "...",
                        "Priority": item_data["priority"],
                        "Status": item_data["status"]
                    })
        
        if attention_items:
            attention_df = pd.DataFrame(attention_items)
            attention_df = attention_df.sort_values(
                by="Priority", 
                key=lambda x: x.map({"Critical": 0, "High": 1, "Medium": 2, "Low": 3})
            )
            st.dataframe(attention_df, use_container_width=True, hide_index=True)
        else:
            st.markdown("""
            <div class="success-card">
                <p>✅ No items requiring immediate attention. All assessed items are compliant.</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Save assessment
        st.session_state.completed_assessments.append({
            "timestamp": datetime.now().isoformat(),
            "system_name": system_name,
            "assessor": assessor_name,
            "overall_score": overall_score,
            "section_scores": section_scores,
            "results": assessment_results
        })

# ============================================
# CASE STUDIES & SCENARIOS
# ============================================
elif current_page == "cases":
    st.markdown('<h1 class="main-header">📊 Case Studies & Scenarios</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-card">
        <p>Learn from real-world scenarios and case studies involving AI ethics in financial services. 
        These examples illustrate common challenges and best practices for responsible AI deployment.</p>
    </div>
    """, unsafe_allow_html=True)
    
    case_tabs = st.tabs([
        "📊 Case Study 1: Credit Scoring",
        "🤖 Case Study 2: Chatbot Failure",
        "📈 Case Study 3: Algorithmic Trading",
        "🎯 Interactive Scenario"
    ])
    
    # Case Study 1: Credit Scoring Bias
    with case_tabs[0]:
        st.markdown("### Case Study: Discriminatory Credit Scoring Model")
        
        st.markdown("""
        <div class="warning-card">
            <h4>⚠️ Scenario Overview</h4>
            <p>A mid-size fintech lender deployed an AI-based credit scoring model that used 
            machine learning to incorporate alternative data sources beyond traditional credit 
            bureau data. After 18 months in production, analysis revealed significant disparate 
            impact against minority applicants.</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        #### 📋 Background
        
        **Company Profile:**
        - Digital lending platform serving underbanked consumers
        - Processing 50,000 loan applications monthly
        - AI model incorporated: social media activity, mobile phone usage patterns, 
          education history, and employment stability
        
        **Initial Goals:**
        - Expand credit access to thin-file consumers
        - Improve prediction accuracy beyond FICO scores
        - Reduce manual underwriting costs
        
        ---
        
        #### ⚠️ What Went Wrong
        
        **Discovery:**
        - Routine fair lending analysis revealed approval rates for minority applicants 
          were 23% lower than for non-minority applicants with similar credit profiles
        - Investigation found several problematic features:
        
        | Feature | Issue | Impact |
        |---------|-------|--------|
        | ZIP code stability | Correlated with neighborhood racial composition | Disparate impact |
        | University attended | Served as proxy for socioeconomic status | Disparate impact |
        | Social media sentiment | Biased training data from predominantly white users | Algorithmic bias |
        | Employment industry | Certain industries had minority overrepresentation | Disparate impact |
        
        **Root Causes:**
        1. Training data reflected historical lending disparities
        2. Feature selection didn't include bias analysis
        3. No ongoing fairness monitoring
        4. Lack of diverse perspectives in model development team
        5. Pressure to maximize approval rates overrode fairness considerations
        
        ---
        
        #### 🔧 Remediation Steps
        
        **Immediate Actions:**
        1. Suspended automated approvals pending review
        2. Notified regulators proactively
        3. Engaged external fair lending experts
        4. Conducted retrospective review of declined applications
        
        **Model Improvements:**
        1. Removed or transformed proxy variables
        2. Implemented adversarial debiasing during training
        3. Added fairness constraints to optimization objective
        4. Introduced mandatory demographic parity thresholds
        
        **Governance Changes:**
        1. Established AI ethics committee with diverse membership
        2. Required pre-deployment fairness testing for all models
        3. Implemented continuous fairness monitoring
        4. Added fair lending expertise to model development team
        
        ---
        
        #### ✅ Key Lessons
        """)
        
        lessons = [
            "**Alternative data requires extra scrutiny** - Novel data sources can introduce unexpected biases",
            "**Fairness testing must be proactive** - Don't wait for complaints or regulatory findings",
            "**Diverse teams catch more issues** - Homogeneous teams have blind spots",
            "**Business pressure doesn't excuse discrimination** - Fairness requirements are non-negotiable",
            "**Continuous monitoring is essential** - Bias can emerge or increase over time"
        ]
        
        for lesson in lessons:
            st.markdown(f"- {lesson}")
        
        st.markdown("""
        ---
        
        #### 📜 Regulatory Implications
        
        | Jurisdiction | Applicable Framework | Potential Consequence |
        |--------------|---------------------|----------------------|
        | United States | ECOA, Fair Housing Act | CFPB enforcement, civil penalties, restitution |
        | European Union | EU AI Act (High-Risk), GDPR | Fines up to €35M or 7% turnover |
        | United Kingdom | Equality Act, FCA Consumer Duty | FCA enforcement action, redress |
        """)
    
    # Case Study 2: Chatbot Failure
    with case_tabs[1]:
        st.markdown("### Case Study: Customer Service Chatbot Failure")
        
        st.markdown("""
        <div class="warning-card">
            <h4>⚠️ Scenario Overview</h4>
            <p>A large retail bank deployed an AI-powered chatbot for customer service. 
            The chatbot provided incorrect information about dispute rights and failed to 
            recognize when customers were invoking legal protections under Regulation E.</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        #### 📋 Background
        
        **Deployment Context:**
        - Major retail bank with 10 million customers
        - Chatbot handled 70% of initial customer contacts
        - Designed to reduce call center volume and costs
        - Built on a large language model with custom fine-tuning
        
        **Intended Capabilities:**
        - Answer account balance and transaction questions
        - Process simple service requests
        - Escalate complex issues to human agents
        
        ---
        
        #### ⚠️ What Went Wrong
        
        **Incident 1: Incorrect Dispute Rights Information**
        - Customer reported unauthorized transaction
        - Chatbot stated customer was "responsible for all transactions made with their card"
        - Failed to mention Regulation E protections limiting liability to $50 if reported within 2 days
        - Customer delayed reporting, incurring higher losses
        
        **Incident 2: Failed Escalation**
        - Customer explicitly stated "I want to dispute this charge under Regulation E"
        - Chatbot did not recognize the regulatory reference
        - Provided generic information about "checking transaction history"
        - No escalation to human agent triggered
        
        **Incident 3: Hallucinated Policy**
        - Customer asked about fee waiver policy
        - Chatbot confidently stated a fee waiver policy that didn't exist
        - Customer recorded conversation and shared on social media
        - Reputational damage and regulatory inquiry followed
        
        **Root Causes:**
        1. LLM not properly fine-tuned on regulatory requirements
        2. Trigger phrases for escalation were too narrow
        3. No content validation against actual bank policies
        4. Insufficient testing with adversarial scenarios
        5. No human review of chatbot training data
        
        ---
        
        #### 🔧 Remediation Steps
        
        **Immediate Actions:**
        1. Added prominent "Speak to Human Agent" option
        2. Disabled responses to dispute-related queries
        3. Notified affected customers of correct dispute rights
        4. Engaged CFPB and state regulators proactively
        
        **Technical Improvements:**
        1. Implemented retrieval-augmented generation (RAG) with policy database
        2. Added rule-based triggers for regulatory terms
        3. Implemented confidence scoring with human handoff below threshold
        4. Created "red team" testing program for adversarial prompts
        
        **Governance Changes:**
        1. Required legal/compliance review of all chatbot responses
        2. Implemented continuous monitoring of customer satisfaction
        3. Created escalation metrics and accountability
        4. Added regular compliance audits of chatbot performance
        
        ---
        
        #### ✅ Key Lessons
        """)
        
        chatbot_lessons = [
            "**Regulatory content requires special handling** - LLMs can hallucinate or misstate legal requirements",
            "**Escalation triggers must be comprehensive** - Customers use varied language to invoke rights",
            "**Confidence thresholds matter** - Low-confidence responses should go to humans",
            "**Continuous monitoring is critical** - Issues may only emerge with diverse user interactions",
            "**Chatbots cannot replace required disclosures** - Legal notices must be accurate and complete"
        ]
        
        for lesson in chatbot_lessons:
            st.markdown(f"- {lesson}")
        
        st.markdown("""
        ---
        
        #### 📜 CFPB Guidance Implications
        
        The CFPB has issued specific guidance on chatbots in consumer finance:
        
        1. **Chatbots cannot replace human interaction** when consumers need substantive help
        2. **Inaccurate information may constitute a UDAAP** (Unfair, Deceptive, or Abusive Act or Practice)
        3. **Must recognize when consumers invoke statutory rights** under Reg E, Reg Z, etc.
        4. **Data privacy and security risks** must be addressed
        5. **Institutions remain responsible** for third-party AI providers
        """)
    
    # Case Study 3: Algorithmic Trading
    with case_tabs[2]:
        st.markdown("### Case Study: AI Trading System Malfunction")
        
        st.markdown("""
        <div class="warning-card">
            <h4>⚠️ Scenario Overview</h4>
            <p>An asset management firm's AI-powered trading system made a series of 
            erratic trades during a period of market volatility, resulting in significant 
            client losses and regulatory scrutiny.</p>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        #### 📋 Background
        
        **System Profile:**
        - Quantitative hedge fund using reinforcement learning for trade execution
        - Managing $2.5 billion in client assets
        - AI system made 85% of trading decisions
        - Trained on 10 years of market data
        
        **Design Intent:**
        - Optimize execution to minimize market impact
        - Identify short-term alpha opportunities
        - Manage portfolio risk in real-time
        
        ---
        
        #### ⚠️ What Went Wrong
        
        **The Incident:**
        - During an unexpected geopolitical event, markets experienced unusual volatility
        - AI system encountered patterns not present in training data
        - System interpreted volatility as alpha opportunity rather than risk
        - Executed a series of leveraged positions that amplified losses
        - Human oversight failed due to speed of execution
        
        **Sequence of Events:**
        
        | Time | Event | Impact |
        |------|-------|--------|
        | 9:31 AM | Market opens with 3% gap down | AI increases position size |
        | 9:45 AM | System adds leverage, misreading signals | Exposure increases 3x |
        | 10:15 AM | Markets stabilize, then drop again | Losses accelerate |
        | 10:30 AM | Kill switch activated | $180M loss crystallized |
        | 11:00 AM | Investigation begins | Trading suspended |
        
        **Root Causes:**
        1. Training data didn't include similar volatility regime
        2. Risk limits were based on historical VaR, not stressed scenarios
        3. Human override was too slow to prevent damage
        4. Lack of circuit breakers for unusual behavior
        5. Overconfidence in AI system after years of success
        
        ---
        
        #### 🔧 Remediation Steps
        
        **Immediate Actions:**
        1. Suspended AI trading pending review
        2. Notified SEC and clients
        3. Engaged independent auditors
        4. Implemented immediate position limits
        
        **System Improvements:**
        1. Added real-time anomaly detection for AI behavior
        2. Implemented regime detection to identify unusual markets
        3. Created hard limits on position size and velocity
        4. Added "uncertainty estimation" to model outputs
        
        **Governance Changes:**
        1. Established AI oversight committee with market risk expertise
        2. Required stress testing under extreme scenarios
        3. Implemented mandatory human approval above thresholds
        4. Created quarterly model review process
        
        ---
        
        #### ✅ Key Lessons
        """)
        
        trading_lessons = [
            "**Historical data has limits** - AI systems can fail in unprecedented conditions",
            "**Speed requires automated safeguards** - Human oversight can't match AI execution speed",
            "**Uncertainty estimation is crucial** - Models should know what they don't know",
            "**Stress testing must go beyond history** - Include hypothetical extreme scenarios",
            "**Success breeds complacency** - Past performance doesn't guarantee future safety"
        ]
        
        for lesson in trading_lessons:
            st.markdown(f"- {lesson}")
    
    # Interactive Scenario
    with case_tabs[3]:
        st.markdown("### 🎯 Interactive Scenario: AI Credit Decision Review")
        
        st.markdown("""
        You are the Chief Risk Officer at a digital lender. Your AI credit scoring team 
        has proposed a new model enhancement. Review the proposal and make a decision.
        """)
        
        st.markdown("""
        ---
        
        #### 📋 Proposal Summary
        
        **Enhancement:** Incorporate social media sentiment analysis into credit scoring
        
        **Claimed Benefits:**
        - 15% improvement in default prediction accuracy
        - Ability to score thin-file applicants
        - Real-time updates to credit assessments
        
        **Data Sources:**
        - Public social media posts
        - Sentiment analysis of text content
        - Network analysis of connections
        - Activity patterns and engagement metrics
        """)
        
        st.markdown("---")
        
        # Decision Points
        st.markdown("#### 🤔 Decision Points")
        
        q1 = st.radio(
            "1. What is your initial reaction to this proposal?",
            [
                "Approve - the accuracy improvement justifies the approach",
                "Conditional approval - needs additional safeguards",
                "Request more information before deciding",
                "Decline - too many risks outweigh benefits"
            ]
        )
        
        q2 = st.multiselect(
            "2. What are your primary concerns? (Select all that apply)",
            [
                "Potential for discrimination based on protected characteristics",
                "Privacy and consent issues",
                "Difficulty in explaining decisions to consumers",
                "Regulatory compliance uncertainty",
                "Data quality and reliability",
                "Reputational risk",
                "Third-party data provider risks"
            ]
        )
        
        q3 = st.multiselect(
            "3. What safeguards would you require before approval? (Select all that apply)",
            [
                "Comprehensive disparate impact testing",
                "Legal review of data collection practices",
                "Consumer disclosure and consent mechanisms",
                "Explainability solution for adverse actions",
                "Ongoing fairness monitoring",
                "Human review for edge cases",
                "Regulatory consultation",
                "External audit of the approach"
            ]
        )
        
        q4 = st.radio(
            "4. Which regulatory frameworks are most relevant to this decision?",
            [
                "ECOA/Fair Credit Reporting Act (US)",
                "EU AI Act and GDPR",
                "UK FCA Consumer Duty",
                "All of the above",
                "None - this is purely a business decision"
            ]
        )
        
        if st.button("Submit Decision"):
            st.markdown("---")
            st.markdown("#### 📊 Analysis of Your Decision")
            
            # Scoring logic
            score = 0
            feedback = []
            
            if "Conditional approval" in q1 or "Request more information" in q1:
                score += 25
                feedback.append("✅ **Good approach:** Requesting conditions or more information shows appropriate caution")
            elif "Decline" in q1:
                score += 20
                feedback.append("⚠️ **Cautious approach:** While declining is safe, conditional approval with strong safeguards may be viable")
            else:
                feedback.append("❌ **Risky approach:** Approving without conditions ignores significant risks")
            
            key_concerns = [
                "Potential for discrimination based on protected characteristics",
                "Privacy and consent issues",
                "Difficulty in explaining decisions to consumers",
                "Regulatory compliance uncertainty"
            ]
            
            matched_concerns = len([c for c in q2 if c in key_concerns])
            score += matched_concerns * 10
            
            if matched_concerns >= 3:
                feedback.append(f"✅ **Strong risk identification:** You identified {matched_concerns}/4 key concerns")
            elif matched_concerns >= 2:
                feedback.append(f"⚠️ **Partial risk identification:** You identified {matched_concerns}/4 key concerns")
            else:
                feedback.append(f"❌ **Risk blind spots:** You only identified {matched_concerns}/4 key concerns")
            
            essential_safeguards = [
                "Comprehensive disparate impact testing",
                "Explainability solution for adverse actions",
                "Ongoing fairness monitoring"
            ]
            
            matched_safeguards = len([s for s in q3 if s in essential_safeguards])
            score += matched_safeguards * 10
            
            if matched_safeguards == 3:
                feedback.append("✅ **Comprehensive safeguards:** You identified all essential safeguards")
            else:
                missing = [s for s in essential_safeguards if s not in q3]
                feedback.append(f"⚠️ **Missing safeguards:** Consider adding: {', '.join(missing)}")
            
            if q4 == "All of the above":
                score += 15
                feedback.append("✅ **Regulatory awareness:** Correct - multiple frameworks apply")
            elif "None" in q4:
                feedback.append("❌ **Regulatory blind spot:** This is definitely a regulated activity")
            
            # Display results
            if score >= 70:
                st.markdown("""
                <div class="success-card">
                    <h4>🎯 Excellent Risk Assessment</h4>
                    <p>Your decision demonstrates strong understanding of AI ethics and regulatory requirements.</p>
                </div>
                """, unsafe_allow_html=True)
            elif score >= 50:
                st.markdown("""
                <div class="warning-card">
                    <h4>⚠️ Good Foundation, Room for Improvement</h4>
                    <p>Your decision shows awareness of key issues but may have some gaps.</p>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="background: #fed7d7; border-left: 4px solid #c53030; padding: 1.5rem; border-radius: 0 8px 8px 0;">
                    <h4>❌ Significant Risk Exposure</h4>
                    <p>Your decision may expose the organization to regulatory and reputational risks.</p>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown("#### Detailed Feedback")
            for item in feedback:
                st.markdown(item)
            
            st.markdown("""
            ---
            
            #### 💡 Expert Analysis
            
            **Recommended Approach:** Conditional approval with comprehensive safeguards
            
            Social media data in credit scoring presents significant risks but may be viable with:
            
            1. **Robust bias testing** - Demographic analysis across all protected classes
            2. **Explainability** - Clear mapping from social signals to credit factors
            3. **Consent** - Transparent disclosure and opt-in mechanisms
            4. **Monitoring** - Ongoing fairness metrics with automatic alerts
            5. **Human oversight** - Review process for declined applications
            6. **Regulatory engagement** - Proactive consultation with CFPB, state regulators
            
            **Key Risk:** Many social media signals correlate with protected characteristics, 
            making discrimination difficult to prevent even with careful design.
            """)

# ============================================
# RESOURCES & DOCUMENTATION
# ============================================
elif current_page == "resources":
    st.markdown('<h1 class="main-header">📚 Resources & Documentation</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-card">
        <p>Access comprehensive resources, official documentation, and learning materials 
        for AI ethics and governance in financial services.</p>
    </div>
    """, unsafe_allow_html=True)
    
    resource_tabs = st.tabs([
        "📜 Official Regulations",
        "📖 Frameworks & Guidelines",
        "🎓 Learning Resources",
        "🔧 Tools & Templates"
    ])
    
    with resource_tabs[0]:
        st.markdown("### Official Regulatory Documents")
        
        st.markdown("""
        #### 🇪🇺 European Union
        
        | Document | Description | Link |
        |----------|-------------|------|
        | EU AI Act (Regulation 2024/1689) | Comprehensive AI regulation | [EUR-Lex](https://eur-lex.europa.eu/legal-content/EN/TXT/?uri=CELEX:32024R1689) |
        | GDPR (Regulation 2016/679) | Data protection framework | [EUR-Lex](https://eur-lex.europa.eu/legal-content/EN/TXT/?uri=CELEX:32016R0679) |
        | AI Office Guidelines | Implementation guidance | [EC Digital Strategy](https://digital-strategy.ec.europa.eu/en/policies/regulatory-framework-ai) |
        
        ---
        
        #### 🇺🇸 United States
        
        | Document | Description | Link |
        |----------|-------------|------|
        | NIST AI RMF 1.0 | AI Risk Management Framework | [NIST](https://www.nist.gov/itl/ai-risk-management-framework) |
        | NIST AI RMF Playbook | Implementation guidance | [NIST Playbook](https://airc.nist.gov/airmf-resources/airmf/) |
        | NIST GenAI Profile | Generative AI guidance | [NIST AI-600-1](https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf) |
        | SR 11-7 | Model Risk Management | [Federal Reserve](https://www.federalreserve.gov/supervisionreg/srletters/sr1107.htm) |
        | CFPB AI Guidance | Consumer protection in AI | [CFPB](https://www.consumerfinance.gov/rules-policy/advanced-technology/) |
        | ECOA | Equal Credit Opportunity Act | [Consumer Finance](https://www.consumerfinance.gov/rules-policy/regulations/1002/) |
        
        ---
        
        #### 🇬🇧 United Kingdom
        
        | Document | Description | Link |
        |----------|-------------|------|
        | FCA AI Update (2024) | FCA approach to AI | [FCA](https://www.fca.org.uk/publication/corporate/ai-update.pdf) |
        | FCA Consumer Duty | Customer outcomes focus | [FCA Consumer Duty](https://www.fca.org.uk/firms/consumer-duty) |
        | AI Survey 2024 | AI usage in UK FS | [Bank of England](https://www.bankofengland.co.uk/report/2024/artificial-intelligence-in-uk-financial-services-2024) |
        | UK Government AI Framework | 5 principles approach | [GOV.UK](https://www.gov.uk/government/publications/ai-regulation-a-pro-innovation-approach) |
        
        ---
        
        #### 🇸🇬 Singapore
        
        | Document | Description | Link |
        |----------|-------------|------|
        | FEAT Principles | Fairness, Ethics, Accountability, Transparency | [MAS](https://www.mas.gov.sg/publications/monographs-or-information-paper/2018/feat) |
        | Veritas Toolkit | FEAT implementation tools | [GitHub](https://github.com/veritas-toolkit/) |
        | MindForge GenAI Whitepaper | GenAI in financial services | [MAS](https://www.mas.gov.sg/schemes-and-initiatives/project-mindforge) |
        """)
    
    with resource_tabs[1]:
        st.markdown("### Frameworks & Industry Guidelines")
        
        st.markdown("""
        #### 🌐 International Standards
        
        - **OECD AI Principles** - International framework for trustworthy AI
        - **ISO/IEC 42001** - AI Management System standard
        - **IEEE 7000 Series** - Ethical considerations in system design
        
        #### 🏦 Financial Services Specific
        
        - **BIS/FSB** - Financial Stability Board AI guidance
        - **IOSCO** - AI in capital markets guidance
        - **IAIS** - Insurance sector AI principles
        
        #### 🔬 Technical Standards
        
        - **NIST SP 800-53** - Security controls (applicable to AI systems)
        - **ISO 27001** - Information security management
        - **SOC 2** - Service organization controls
        
        #### 🎯 Industry Best Practices
        
        | Organization | Resource | Focus |
        |--------------|----------|-------|
        | Partnership on AI | Best Practices | Cross-industry AI ethics |
        | IEEE | Ethically Aligned Design | Technical ethics standards |
        | WEF | AI Governance Alliance | Multi-stakeholder governance |
        | GFIN | Regulatory Innovation | Cross-border fintech regulation |
        """)
    
    with resource_tabs[2]:
        st.markdown("### Learning Resources")
        
        st.markdown("""
        #### 📚 Recommended Reading
        
        **Books:**
        - "Weapons of Math Destruction" - Cathy O'Neil
        - "Algorithms of Oppression" - Safiya Noble
        - "The Alignment Problem" - Brian Christian
        - "Artificial Intelligence: A Guide for Thinking Humans" - Melanie Mitchell
        
        **Academic Papers:**
        - "Fairness and Machine Learning" - Barocas, Hardt, Narayanan
        - "Algorithmic Discrimination in Credit Domain" - Springer AI & Society
        - "Explainable AI in Financial Services" - Various authors
        
        #### 🎓 Online Courses
        
        | Course | Provider | Focus |
        |--------|----------|-------|
        | AI Ethics | MIT | Foundational ethics |
        | Responsible AI | Google | Practical implementation |
        | Fairness in ML | Microsoft | Technical fairness |
        | AI Governance | Stanford HAI | Governance frameworks |
        
        #### 🎥 Video Resources
        
        - NIST AI RMF Explainer Video
        - FCA AI Lab Webinars
        - MAS FinTech Festival Sessions
        - CFPB Consumer Protection Webinars
        
        #### 🎪 Conferences & Events
        
        - Singapore FinTech Festival
        - Money 20/20
        - AI & Big Data Expo
        - RegTech Summit
        """)
    
    with resource_tabs[3]:
        st.markdown("### Tools & Templates")
        
        st.markdown("""
        #### 🔧 Assessment Tools
        
        | Tool | Purpose | Source |
        |------|---------|--------|
        | Veritas Toolkit | FEAT assessment | MAS/GitHub |
        | AI Fairness 360 | Bias detection | IBM |
        | Fairlearn | Fairness assessment | Microsoft |
        | What-If Tool | Model exploration | Google |
        | SHAP | Explainability | Open source |
        | LIME | Local explanations | Open source |
        
        #### 📋 Templates
        
        This toolkit provides the following templates (available in Export section):
        
        - **AI Risk Assessment Template** - Comprehensive risk identification
        - **Governance Framework Template** - Organizational structure and policies
        - **Ethical Assessment Checklist** - Pre-deployment review
        - **Incident Response Template** - AI failure handling
        - **Regulatory Mapping Template** - Compliance tracking
        
        #### 🖥️ Technical Resources
        
        **Explainability Libraries:**
        - SHAP (SHapley Additive exPlanations)
        - LIME (Local Interpretable Model-agnostic Explanations)
        - Alibi
        - Captum (PyTorch)
        
        **Fairness Libraries:**
        - Fairlearn (Microsoft)
        - AI Fairness 360 (IBM)
        - Aequitas
        - Themis-ML
        
        **Model Monitoring:**
        - Evidently AI
        - Fiddler AI
        - Arize AI
        - WhyLabs
        """)

# ============================================
# EXPORT & REPORTS
# ============================================
elif current_page == "export":
    st.markdown('<h1 class="main-header">💾 Export & Reports</h1>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="info-card">
        <p>Generate and export reports from your assessments for documentation, 
        audit, and compliance purposes.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Available Exports
    st.markdown("### 📤 Available Exports")
    
    export_col1, export_col2 = st.columns(2)
    
    with export_col1:
        st.markdown("#### 📊 Risk Assessment Report")
        if st.session_state.risk_assessment:
            st.markdown("""
            <div class="success-card">
                <p>✅ Risk assessment data available</p>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("📥 Export Risk Assessment (JSON)"):
                risk_json = json.dumps(st.session_state.risk_assessment, indent=2, default=str)
                st.download_button(
                    label="Download Risk Assessment",
                    data=risk_json,
                    file_name=f"risk_assessment_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
        else:
            st.markdown("""
            <div class="warning-card">
                <p>⚠️ No risk assessment completed. Complete the Risk Identification Tool first.</p>
            </div>
            """, unsafe_allow_html=True)
    
    with export_col2:
        st.markdown("#### 🔧 Governance Framework")
        if st.session_state.governance_plan:
            st.markdown("""
            <div class="success-card">
                <p>✅ Governance framework data available</p>
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("📥 Export Governance Framework (JSON)"):
                gov_json = json.dumps(st.session_state.governance_plan, indent=2, default=str)
                st.download_button(
                    label="Download Governance Framework",
                    data=gov_json,
                    file_name=f"governance_framework_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )
        else:
            st.markdown("""
            <div class="warning-card">
                <p>⚠️ No governance framework created. Use the Governance Framework Builder first.</p>
            </div>
            """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Completed Assessments
    st.markdown("### ✅ Completed Ethical Assessments")
    
    if st.session_state.completed_assessments:
        for i, assessment in enumerate(st.session_state.completed_assessments):
            with st.expander(f"Assessment: {assessment.get('system_name', 'Unnamed')} - {assessment['timestamp'][:10]}"):
                st.json(assessment)
                
                assessment_json = json.dumps(assessment, indent=2, default=str)
                st.download_button(
                    label=f"📥 Download Assessment {i+1}",
                    data=assessment_json,
                    file_name=f"ethical_assessment_{i+1}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    key=f"download_assessment_{i}"
                )
    else:
        st.markdown("""
        <div class="warning-card">
            <p>⚠️ No ethical assessments completed. Use the Ethical Assessment Checklist first.</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
# ... inside elif current_page == "export": ...

    # Comprehensive Report Generator
    st.markdown("### 📋 Generate Comprehensive Report")
    
    st.markdown("""
    Generate a comprehensive report combining all available assessment data.
    """)
    
    # Selection for report format
    report_format = st.selectbox(
        "Report Format:",
        ["DOCX (Microsoft Word)", "JSON (Machine-readable)"]
    )
    
    if st.button("📄 Generate Comprehensive Report", type="primary"):
        
        # JSON Logic (kept for machine readability)
        if report_format == "JSON (Machine-readable)":
            report_data = {
                "report_metadata": {
                    "generated_at": datetime.now().isoformat(),
                    "tool_version": "1.0.0"
                },
                "risk_assessment": st.session_state.risk_assessment,
                "governance_framework": st.session_state.governance_plan,
                "ethical_assessments": st.session_state.completed_assessments
            }
            report_json = json.dumps(report_data, indent=2, default=str)
            st.download_button(
                label="📥 Download Comprehensive Report (JSON)",
                data=report_json,
                file_name=f"ai_ethics_report_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json"
            )

        # DOCX Logic (New implementation)
        else:
            # Initialize Document
            doc = Document()
            
            # Title
            title = doc.add_heading('FinTech AI Ethics & Governance Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
            doc.add_paragraph("Tool Version: 1.0.0")
            doc.add_paragraph("_" * 50) # Horizontal line simulation

            # 1. Executive Summary
            doc.add_heading('1. Executive Summary', level=1)
            doc.add_paragraph(
                "This report summarizes the AI ethics and governance assessment conducted using "
                "the FinTech AI Ethics & Governance Toolkit. It includes risk identification, "
                "governance framework definitions, and ethical assessment checklists."
            )

            # 2. Risk Assessment Section
            doc.add_heading('2. Risk Assessment', level=1)
            if st.session_state.risk_assessment:
                risk_data = st.session_state.risk_assessment
                p = doc.add_paragraph()
                p.add_run("Use Case: ").bold = True
                p.add_run(f"{risk_data.get('use_case', 'N/A')}\n")
                p.add_run("Risk Level: ").bold = True
                p.add_run(f"{risk_data.get('risk_level', 'Unknown')}\n")
                p.add_run("Overall Score: ").bold = True
                p.add_run(f"{risk_data.get('overall_score', 0):.1f}/100")
                
                doc.add_heading('Category Scores:', level=2)
                for cat, score in risk_data.get('category_scores', {}).items():
                    doc.add_paragraph(f"{cat}: {score:.1f}%", style='List Bullet')
            else:
                doc.add_paragraph("Risk assessment not yet completed.")

            # 3. Governance Framework Section
            doc.add_heading('3. Governance Framework', level=1)
            if st.session_state.governance_plan:
                gov_data = st.session_state.governance_plan
                
                doc.add_heading('Key Roles', level=2)
                structure = gov_data.get('structure', {})
                if structure:
                    doc.add_paragraph(f"AI Officer: {structure.get('ai_officer', 'Not defined')}")
                    doc.add_paragraph(f"Risk Owner: {structure.get('ai_risk_owner', 'Not defined')}")
                
                doc.add_heading('Policy Status', level=2)
                for policy, status in gov_data.get('policies', {}).items():
                    doc.add_paragraph(f"{policy}: {status}", style='List Bullet')
            else:
                doc.add_paragraph("Governance framework not yet generated.")

            # 4. Ethical Assessments Section
            doc.add_heading('4. Ethical Assessments', level=1)
            if st.session_state.completed_assessments:
                for i, assessment in enumerate(st.session_state.completed_assessments, 1):
                    doc.add_heading(f"Assessment {i}: {assessment.get('system_name')}", level=2)
                    doc.add_paragraph(f"Assessor: {assessment.get('assessor')}")
                    doc.add_paragraph(f"Date: {assessment.get('timestamp')}")
                    doc.add_paragraph(f"Score: {assessment.get('overall_score'):.1f}%")
                    
                    # Add section scores table
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Section'
                    hdr_cells[1].text = 'Compliance Score'
                    
                    for section, score in assessment.get('section_scores', {}).items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(section)
                        row_cells[1].text = f"{score:.1f}%"
                    doc.add_paragraph() # Spacer
            else:
                doc.add_paragraph("No ethical assessments completed.")

            # Save to Memory Buffer
            bio = io.BytesIO()
            doc.save(bio)
            
            st.download_button(
                label="📥 Download Comprehensive Report (DOCX)",
                data=bio.getvalue(),
                file_name=f"ai_ethics_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        st.success("Report generated successfully!")

# Footer
st.markdown("---")
st.markdown("""
<div class="footer">
    <h4>FinTech AI Ethics & Governance Toolkit</h4>
    <p>Version 1.0.0 | Last Updated: November 2025</p>
    <p>Developed by Professor Vangelis Tsiligkiris</p>
    <p>Built for educational and professional use in financial services AI governance.</p>
    <p style="font-size: 0.8rem; margin-top: 1rem;">
        This tool is for educational purposes. Always consult with legal and compliance 
        professionals for specific regulatory requirements.
    </p>
</div>
""", unsafe_allow_html=True)
